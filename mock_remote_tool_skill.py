from fastapi import FastAPI, Header
from pydantic import BaseModel
from typing import Any, Dict, List, Optional, Tuple
from pathlib import Path
import concurrent.futures
import json
import logging
import re
import time

logger = logging.getLogger(__name__)

try:
    from src.llm.llm import get_llm_by_type
except Exception:  # pragma: no cover
    def get_llm_by_type(*_args, **_kwargs):  # type: ignore
        raise RuntimeError("LLM dependencies are not available")

app = FastAPI()


class ToolRequest(BaseModel):
    tool: str
    arguments: Dict[str, Any]


class SkillRequest(BaseModel):
    skill: str
    arguments: Dict[str, Any]


_SAMPLE_CACHE: Optional[Dict[str, Any]] = None
_TODO_CACHE: Optional[Dict[str, Any]] = None
_UNICORN_CACHE: Optional[Dict[str, Any]] = None
_RISK_CACHE: Optional[Dict[str, Any]] = None
_EMAIL_CACHE: Optional[Dict[str, Any]] = None
_SCHEDULE_CACHE: Optional[Dict[str, Any]] = None
_KNOWLEDGE_CACHE: Optional[Dict[str, Any]] = None
_SALARY_CACHE: Optional[Dict[str, Any]] = None
_TEMPLATE_CACHE: Optional[Dict[str, Any]] = None


def _read_json(path: Path) -> Dict[str, Any]:
    text = path.read_text(encoding="utf-8-sig")
    if text and text[0] == "\ufeff":
        text = text.lstrip("\ufeff")
    return json.loads(text)


def _invoke_with_timeout(func, timeout_sec: float):
    with concurrent.futures.ThreadPoolExecutor(max_workers=1) as executor:
        future = executor.submit(func)
        return future.result(timeout=timeout_sec)


def _sample_path() -> Path:
    return Path(__file__).resolve().parent / "assets" / "person_info_sample.json"


def _load_sample() -> Dict[str, Any]:
    global _SAMPLE_CACHE
    if _SAMPLE_CACHE is not None:
        return _SAMPLE_CACHE
    path = _sample_path()
    if not path.exists():
        raise FileNotFoundError(f"Sample data not found: {path}")
    _SAMPLE_CACHE = _read_json(path)
    return _SAMPLE_CACHE


def _todo_path() -> Path:
    return Path(__file__).resolve().parent / "assets" / "todo_sample.json"


def _load_todos() -> Dict[str, Any]:
    global _TODO_CACHE
    if _TODO_CACHE is not None:
        return _TODO_CACHE
    path = _todo_path()
    if not path.exists():
        raise FileNotFoundError(f"Todo data not found: {path}")
    _TODO_CACHE = _read_json(path)
    return _TODO_CACHE


def _unicorn_path() -> Path:
    return Path(__file__).resolve().parent / "assets" / "unicorn_db.json"


def _load_unicorns() -> Dict[str, Any]:
    global _UNICORN_CACHE
    if _UNICORN_CACHE is not None:
        return _UNICORN_CACHE
    path = _unicorn_path()
    if not path.exists():
        raise FileNotFoundError(f"Unicorn DB not found: {path}")
    _UNICORN_CACHE = _read_json(path)
    return _UNICORN_CACHE


def _risk_path() -> Path:
    return Path(__file__).resolve().parent / "assets" / "credit_risk_db.json"


def _load_risks() -> Dict[str, Any]:
    global _RISK_CACHE
    if _RISK_CACHE is not None:
        return _RISK_CACHE
    path = _risk_path()
    if not path.exists():
        raise FileNotFoundError(f"Risk DB not found: {path}")
    _RISK_CACHE = _read_json(path)
    return _RISK_CACHE


def _email_path() -> Path:
    return Path(__file__).resolve().parent / "assets" / "email_log.json"


def _load_emails() -> Dict[str, Any]:
    global _EMAIL_CACHE
    if _EMAIL_CACHE is not None:
        return _EMAIL_CACHE
    path = _email_path()
    if not path.exists():
        raise FileNotFoundError(f"Email log not found: {path}")
    _EMAIL_CACHE = _read_json(path)
    return _EMAIL_CACHE


def _schedule_path() -> Path:
    return Path(__file__).resolve().parent / "assets" / "visit_schedule.json"


def _load_schedules() -> Dict[str, Any]:
    global _SCHEDULE_CACHE
    if _SCHEDULE_CACHE is not None:
        return _SCHEDULE_CACHE
    path = _schedule_path()
    if not path.exists():
        raise FileNotFoundError(f"Schedule DB not found: {path}")
    _SCHEDULE_CACHE = _read_json(path)
    return _SCHEDULE_CACHE


def _salary_path() -> Path:
    return Path(__file__).resolve().parent / "assets" / "mock_salary_db.json"


def _load_salary() -> Dict[str, Any]:
    global _SALARY_CACHE
    if _SALARY_CACHE is not None:
        return _SALARY_CACHE
    path = _salary_path()
    if not path.exists():
        raise FileNotFoundError(f"Salary DB not found: {path}")
    _SALARY_CACHE = _read_json(path)
    return _SALARY_CACHE


def _template_path() -> Path:
    return Path(__file__).resolve().parent / "assets" / "document_templates.json"


def _load_templates() -> Dict[str, Any]:
    global _TEMPLATE_CACHE
    if _TEMPLATE_CACHE is not None:
        return _TEMPLATE_CACHE
    path = _template_path()
    if not path.exists():
        raise FileNotFoundError(f"Template file not found: {path}")
    _TEMPLATE_CACHE = _read_json(path)
    return _TEMPLATE_CACHE


def _number_to_chinese(num: float) -> str:
    """Convert number to Chinese uppercase format"""
    if num == 0:
        return "零元整"

    units = ["", "拾", "佰", "仟", "万", "拾", "佰", "仟", "亿"]
    digits = ["零", "壹", "贰", "叁", "肆", "伍", "陆", "柒", "捌", "玖"]

    # Split into integer and decimal parts
    int_part = int(num)
    decimal_part = round((num - int_part) * 100)

    # Convert integer part
    result = ""
    str_num = str(int_part)
    length = len(str_num)

    for i, digit in enumerate(str_num):
        digit_val = int(digit)
        if digit_val != 0:
            result += digits[digit_val] + units[length - i - 1]
        elif result and result[-1] != "零":
            result += "零"

    result = result.rstrip("零") + "元"

    # Add decimal part
    if decimal_part > 0:
        jiao = decimal_part // 10
        fen = decimal_part % 10
        if jiao > 0:
            result += digits[jiao] + "角"
        if fen > 0:
            result += digits[fen] + "分"
    else:
        result += "整"

    return result


def _knowledge_path() -> Path:
    return Path(__file__).resolve().parent / "assets" / "knowledge_base.json"


def _load_knowledge() -> Dict[str, Any]:
    global _KNOWLEDGE_CACHE
    if _KNOWLEDGE_CACHE is not None:
        return _KNOWLEDGE_CACHE
    path = _knowledge_path()
    if not path.exists():
        raise FileNotFoundError(f"Knowledge base not found: {path}")
    _KNOWLEDGE_CACHE = _read_json(path)
    return _KNOWLEDGE_CACHE


def _flatten_text(person: Dict[str, Any]) -> str:
    parts: List[str] = []
    for value in person.values():
        if isinstance(value, str):
            parts.append(value)
        elif isinstance(value, (int, float)):
            parts.append(str(value))
    return " ".join(parts)


def _extract_year_numbers(*values: Optional[str]) -> List[float]:
    numbers: List[float] = []
    pattern = re.compile(r"(\d+(?:\.\d+)?)\s*年")
    for value in values:
        if not value:
            continue
        for match in pattern.findall(value):
            try:
                numbers.append(float(match))
            except ValueError:
                continue
    return numbers


def _normalize_list(value: Any) -> List[str]:
    if value is None:
        return []
    if isinstance(value, list):
        return [str(v) for v in value if v is not None]
    return [str(value)]


def _matches_condition(person: Dict[str, Any], condition: Dict[str, Any]) -> bool:
    name = str(condition.get("cndName") or condition.get("name") or "").strip()
    values = _normalize_list(condition.get("cndValList") or condition.get("values"))
    range_values = condition.get("rangeValueList") or condition.get("range_values") or []

    text = _flatten_text(person)
    value_ok = True
    if values:
        value_ok = any(v in text for v in values)

    range_ok = True
    if range_values:
        range_ok = False
        years = _extract_year_numbers(
            person.get("attr1DefDsc"),
            person.get("attr2DefDsc"),
            person.get("attr3DefDsc"),
            person.get("attr4DefDsc"),
        )
        for item in range_values:
            start_val = item.get("startVal")
            if start_val is None:
                continue
            try:
                start_num = float(start_val)
            except ValueError:
                continue
            if any(year >= start_num for year in years):
                range_ok = True
                break

    if name in {"在职状态"}:
        status = person.get("empeStdsc")
        value_ok = not values or any(v in str(status) for v in values)

    if name in {"机构名称"}:
        inst = f"{person.get('holdposInstNm') or ''} {person.get('instFullNm') or ''}"
        value_ok = not values or any(v in inst for v in values)

    if name in {"姓名", "人员姓名"}:
        name_value = person.get("adtEmpeNm") or ""
        value_ok = not values or any(v in name_value for v in values)

    return value_ok and range_ok


def _get_birth_year(person: Dict[str, Any]) -> Optional[int]:
    brth = person.get("brthDt")
    if not brth:
        return None
    try:
        return int(str(brth)[:4])
    except Exception:
        return None


def _get_age(person: Dict[str, Any]) -> Optional[int]:
    age_val = person.get("age")
    if age_val is None:
        return None
    try:
        return int(str(age_val))
    except Exception:
        return None


def _get_float_value(value: Any) -> Optional[float]:
    if value is None:
        return None
    try:
        return float(str(value))
    except Exception:
        return None


def _get_experience_years(person: Dict[str, Any]) -> Optional[float]:
    years = _extract_year_numbers(
        person.get("attr1DefDsc"),
        person.get("attr2DefDsc"),
        person.get("attr3DefDsc"),
        person.get("attr4DefDsc"),
    )
    if years:
        return max(years)
    return None


def _get_base_years(person: Dict[str, Any]) -> Optional[float]:
    values = []
    for key in ("rmrk1Inf", "rmrk2Inf"):
        val = person.get(key)
        num = _get_float_value(val)
        if num is not None:
            values.append(num)
        values.extend(_extract_year_numbers(str(val) if val is not None else ""))
    if values:
        return max(values)
    return None


def _filter_people(
    people: List[Dict[str, Any]],
    conditions: List[Dict[str, Any]],
    keyword: Optional[str],
    keywords: Optional[List[str]],
    gender: Optional[str],
    birth_year_range: Optional[Tuple[int, int]],
    age_min: Optional[int],
    age_max: Optional[int],
    job_keywords: Optional[List[str]],
    org_keywords: Optional[List[str]],
    education_keywords: Optional[List[str]],
    education_any_keywords: Optional[List[str]],
    experience_min: Optional[float],
    experience_max: Optional[float],
    work_years_min: Optional[float],
    work_years_max: Optional[float],
    bank_years_min: Optional[float],
    bank_years_max: Optional[float],
    base_years_min: Optional[float],
    base_years_max: Optional[float],
) -> List[Dict[str, Any]]:
    if keyword:
        conditions = list(conditions)
        conditions.append({"cndName": "关键词", "cndValList": [keyword]})

    if keywords:
        conditions = list(conditions)
        conditions.append({"cndName": "关键词", "cndValList": list(keywords)})

    if not conditions:
        conditions = []

    results: List[Dict[str, Any]] = []
    for person in people:
        if conditions and not all(_matches_condition(person, cond) for cond in conditions):
            continue

        if gender:
            if gender not in str(person.get("gnd") or ""):
                continue

        if birth_year_range:
            birth_year = _get_birth_year(person)
            if birth_year is None:
                continue
            start, end = birth_year_range
            if not (start <= birth_year <= end):
                continue

        if age_min is not None or age_max is not None:
            age_val = _get_age(person)
            if age_val is None:
                continue
            if age_min is not None and age_val < age_min:
                continue
            if age_max is not None and age_val > age_max:
                continue

        if job_keywords:
            job_text = " ".join(
                [
                    str(person.get("nwgntPstNm") or ""),
                    str(person.get("tcoPostNm") or ""),
                    str(person.get("postCmnt") or ""),
                    str(person.get("seqNm") or ""),
                    str(person.get("pstCtlg") or ""),
                ]
            )
            # 改为"任意包含"：只要有一个关键词匹配即可
            if not any(k in job_text for k in job_keywords):
                continue

        if org_keywords:
            org_text = " ".join(
                [
                    str(person.get("holdposInstNm") or ""),
                    str(person.get("instFullNm") or ""),
                    str(person.get("boFullnm") or ""),
                    str(person.get("instAttrChnNm") or ""),
                    str(person.get("mngbkInstLvlNm") or ""),  # 添加机构级别字段
                ]
            )
            # 改为"任意包含"：只要有一个关键词匹配即可
            if not any(k in org_text for k in org_keywords):
                continue

        if education_any_keywords:
            edu_text = " ".join(
                [
                    str(person.get("education") or ""),
                    str(person.get("hgstEddgrNm") or ""),
                    str(person.get("grdtUnvrstNm") or ""),
                    str(person.get("shlNm") or ""),
                    str(person.get("mjrNm") or ""),
                ]
            )
            if not any(k in edu_text for k in education_any_keywords):
                continue

        if education_keywords:
            edu_text = " ".join(
                [
                    str(person.get("education") or ""),
                    str(person.get("hgstEddgrNm") or ""),
                    str(person.get("grdtUnvrstNm") or ""),
                    str(person.get("shlNm") or ""),
                    str(person.get("mjrNm") or ""),
                ]
            )
            if not all(k in edu_text for k in education_keywords):
                continue

        if experience_min is not None or experience_max is not None:
            exp_years = _get_experience_years(person)
            if exp_years is None:
                continue
            if experience_min is not None and exp_years < experience_min:
                continue
            if experience_max is not None and exp_years > experience_max:
                continue

        if work_years_min is not None or work_years_max is not None:
            work_years = _get_float_value(person.get("pcsTrdYrlmt"))
            if work_years is None:
                continue
            if work_years_min is not None and work_years < work_years_min:
                continue
            if work_years_max is not None and work_years > work_years_max:
                continue

        if bank_years_min is not None or bank_years_max is not None:
            bank_years = _get_float_value(person.get("mbshYrlmt"))
            if bank_years is None:
                continue
            if bank_years_min is not None and bank_years < bank_years_min:
                continue
            if bank_years_max is not None and bank_years > bank_years_max:
                continue

        if base_years_min is not None or base_years_max is not None:
            base_years = _get_base_years(person)
            if base_years is None:
                continue
            if base_years_min is not None and base_years < base_years_min:
                continue
            if base_years_max is not None and base_years > base_years_max:
                continue

        results.append(person)

    return results


def _build_request_xml(condition_list: List[Dict[str, Any]]) -> str:
    payload = {
        "tx26IdxVal": "04",
        "sessionId": str(int(time.time() * 1000)),
        "source": "web",
        "conditionList": condition_list,
    }
    data = json.dumps(payload, ensure_ascii=False)
    return (
        "<ENTITY>\n"
        "  <Mnplt_TpCd><![CDATA[]]></Mnplt_TpCd>\n"
        "  <Tx_26_Idx_Val><![CDATA[]]></Tx_26_Idx_Val>\n"
        f"  <Data_Stc_Dsc><![CDATA[{data}]]></Data_Stc_Dsc>\n"
        "</ENTITY>"
    )


def _build_response_xml(matched_count: int) -> str:
    summary = {"matched_count": matched_count}
    message_content = "```json\n" + json.dumps(summary, ensure_ascii=False) + "\n```"
    payload = {
        "choices": [
            {
                "finish_reason": "stop",
                "index": 0,
                "message": {"content": message_content, "role": "assistant", "tool_calls": []},
            }
        ],
        "created": int(time.time() * 1000),
        "notes": "提示：接口传入列表！",
        "traceId": f"trace_{int(time.time() * 1000)}",
        "usage": {"completion_tokens": 0, "prompt_tokens": 0, "total_tokens": 0},
    }
    data = json.dumps(payload, ensure_ascii=False)
    return (
        "<ENTITY>\n"
        f"<Data_Enqr_Rslt><![CDATA[{data}]]></Data_Enqr_Rslt>\n"
        "<codeid><![CDATA[20000]]></codeid>\n"
        "</ENTITY>"
    )


def _build_markdown_report(title: str, sections: List[Dict[str, Any]]) -> str:
    lines = [f"# {title}", ""]
    for sec in sections:
        if not isinstance(sec, dict):
            continue
        heading = sec.get("heading", "")
        content = sec.get("content", "")
        if heading:
            lines.append(f"## {heading}")
        if content:
            lines.append(str(content))
        lines.append("")
    return "\n".join(lines).strip()


@app.get("/health")
async def health():
    return {"status": "ok"}


@app.post("/tool")
async def tool(req: ToolRequest, authorization: Optional[str] = Header(default=None)):
    global _TODO_CACHE, _EMAIL_CACHE, _SCHEDULE_CACHE
    if req.tool == "remote_weather_tool":
        location = req.arguments.get("location", "")
        result = f"[remote-tool] {location} weather: sunny 20C"
    elif req.tool == "remote_person_info_tool":
        try:
            sample = _load_sample()
            conditions = req.arguments.get("conditions") or req.arguments.get("conditionList") or []
            keyword = req.arguments.get("keyword")
            keywords = req.arguments.get("keywords")
            gender = req.arguments.get("gender")
            birth_range = req.arguments.get("birth_year_range")
            age_min = req.arguments.get("age_min")
            age_max = req.arguments.get("age_max")
            job_keywords = req.arguments.get("job_keywords")
            org_keywords = req.arguments.get("org_keywords")
            education_keywords = req.arguments.get("education_keywords")
            education_any_keywords = req.arguments.get("education_any_keywords")
            experience_min = _get_float_value(req.arguments.get("experience_min"))
            experience_max = _get_float_value(req.arguments.get("experience_max"))
            work_years_min = _get_float_value(req.arguments.get("work_years_min"))
            work_years_max = _get_float_value(req.arguments.get("work_years_max"))
            bank_years_min = _get_float_value(req.arguments.get("bank_years_min"))
            bank_years_max = _get_float_value(req.arguments.get("bank_years_max"))
            base_years_min = _get_float_value(req.arguments.get("base_years_min"))
            base_years_max = _get_float_value(req.arguments.get("base_years_max"))

            # 记录提取的参数
            logger.info(f"Person query parameters:")
            logger.info(f"  - job_keywords: {job_keywords}")
            logger.info(f"  - org_keywords: {org_keywords}")
            logger.info(f"  - birth_year_range: {birth_range}")
            logger.info(f"  - gender: {gender}")
            logger.info(f"  - keyword: {keyword}")

            birth_year_range: Optional[Tuple[int, int]] = None
            if isinstance(birth_range, (list, tuple)) and len(birth_range) == 2:
                try:
                    birth_year_range = (int(birth_range[0]), int(birth_range[1]))
                except Exception:
                    birth_year_range = None

            people = sample.get("personInfoList", [])
            logger.info(f"Total people in database: {len(people)}")

            filtered = _filter_people(
                people,
                conditions,
                keyword,
                keywords,
                gender,
                birth_year_range,
                age_min,
                age_max,
                job_keywords,
                org_keywords,
                education_keywords,
                education_any_keywords,
                experience_min,
                experience_max,
                work_years_min,
                work_years_max,
                bank_years_min,
                bank_years_max,
                base_years_min,
                base_years_max,
            )

            logger.info(f"Filtered results: {len(filtered)} people matched")
            if len(filtered) > 0:
                logger.info(f"First matched person: {filtered[0].get('adtEmpeNm', 'N/A')}")

            limit = req.arguments.get("limit")
            if isinstance(limit, int) and limit > 0:
                filtered = filtered[:limit]

            payload: Dict[str, Any] = {
                "status": "success",
                "matched_count": len(filtered),
                "condition_list": conditions,
            }
            if req.arguments.get("return_xml", True):
                payload["request_xml"] = _build_request_xml(conditions)
                payload["response_xml"] = _build_response_xml(len(filtered))
            if req.arguments.get("return_detail", True):
                payload["detail"] = {
                    "authPersonPropertyMap": sample.get("authPersonPropertyMap", {}),
                    "personInfoList": filtered,
                }
            result = payload
        except Exception as exc:
            result = {"status": "error", "error": str(exc)}
    elif req.tool == "remote_todo_query_tool":
        try:
            data = _load_todos()
            todos = data.get("todos", [])
            action = req.arguments.get("action", "read")

            if action == "create":
                payload = req.arguments.get("todo") or {}
                if not isinstance(payload, dict):
                    raise ValueError("todo must be an object")
                new_id = payload.get("id") or f"todo-{len(todos)+1:03d}"
                payload["id"] = new_id
                todos.append(payload)
                data["todos"] = todos
                _TODO_CACHE = data
                _todo_path().write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
                result = {"status": "success", "created_id": new_id, "todo": payload}
            elif action == "update":
                todo_id = req.arguments.get("id")
                updates = req.arguments.get("updates") or {}
                if not todo_id:
                    raise ValueError("id is required for update")
                updated = None
                for item in todos:
                    if item.get("id") == todo_id:
                        if isinstance(updates, dict):
                            item.update(updates)
                        updated = item
                        break
                if updated is None:
                    raise ValueError("todo not found")
                data["todos"] = todos
                _TODO_CACHE = data
                _todo_path().write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
                result = {"status": "success", "updated": updated}
            elif action == "delete":
                todo_id = req.arguments.get("id")
                if not todo_id:
                    raise ValueError("id is required for delete")
                before = len(todos)
                todos = [t for t in todos if t.get("id") != todo_id]
                if len(todos) == before:
                    raise ValueError("todo not found")
                data["todos"] = todos
                _TODO_CACHE = data
                _todo_path().write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
                result = {"status": "success", "deleted_id": todo_id}
            else:
                start_date = req.arguments.get("start_date")
                end_date = req.arguments.get("end_date")
                status = req.arguments.get("status")
                user_id = req.arguments.get("user_id")
                employee_id = req.arguments.get("employee_id")

                filtered = []
                for item in todos:
                    due = item.get("due_date")
                    if start_date and due and due < start_date:
                        continue
                    if end_date and due and due > end_date:
                        continue
                    if status and str(item.get("status")) != str(status):
                        continue
                    if user_id and str(item.get("user_id")) != str(user_id):
                        continue
                    if employee_id and str(item.get("employee_id")) != str(employee_id):
                        continue
                    filtered.append(item)

                result = {
                    "status": "success",
                    "matched_count": len(filtered),
                    "todos": filtered,
                    "range": {"start_date": start_date, "end_date": end_date},
                }
        except Exception as exc:
            result = {"status": "error", "error": str(exc)}
    elif req.tool == "remote_unicorn_db_tool":
        try:
            data = _load_unicorns()
            records = data.get("records", [])
            limit = req.arguments.get("limit", 5)
            industry = req.arguments.get("industry")
            status = req.arguments.get("status", "独角兽")

            filtered = []
            for item in records:
                if status and item.get("status") != status:
                    continue
                if industry and item.get("industry") != industry:
                    continue
                filtered.append(item)

            result = {
                "status": "success",
                "matched_count": len(filtered[:limit]),
                "records": filtered[:limit],
                "schema": data.get("schema"),
            }
        except Exception as exc:
            result = {"status": "error", "error": str(exc)}
    elif req.tool == "remote_credit_risk_db_tool":
        try:
            data = _load_risks()
            records = data.get("records", [])
            company_ids = req.arguments.get("company_ids") or []
            if isinstance(company_ids, str):
                company_ids = [company_ids]
            if company_ids:
                records = [r for r in records if r.get("company_id") in company_ids]
            result = {
                "status": "success",
                "matched_count": len(records),
                "records": records,
                "schema": data.get("schema"),
            }
        except Exception as exc:
            result = {"status": "error", "error": str(exc)}
    elif req.tool == "remote_report_builder_tool":
        try:
            # 支持两种输入格式：
            # 1. 结构化格式：{"title": "...", "sections": [...]}
            # 2. 原始数据格式：{"data": [...]} - 需要 LLM 生成报告

            data = req.arguments.get("data")
            title = req.arguments.get("title", "分析报告")
            sections = req.arguments.get("sections") or []
            instruction = req.arguments.get("instruction") or "生成清晰、专业的Markdown分析报告。"
            llm_timeout_sec = req.arguments.get("llm_timeout_sec", 60)  # 增加超时时间

            # 如果提供了原始数据，使用 LLM 生成完整报告
            if data is not None:
                try:
                    llm = get_llm_by_type("basic")

                    # 构建专业的报告生成 prompt
                    prompt = f"""你是一位专业的企业分析报告撰写专家。请基于以下数据生成一份专业的Markdown格式分析报告。

# 报告要求

1. 报告结构：
   - 标题和概述
   - 核心发现
   - 详细分析（按企业分组）
   - 风险评估总结
   - 结论与建议

2. 写作风格：
   - 使用专业、客观的语言
   - 数据驱动，有理有据
   - 突出关键指标和风险点
   - 使用表格和列表增强可读性

3. 分析维度：
   - 信用评分分布
   - 风险等级分类
   - 现金流健康状况
   - 债务比率分析
   - 违约事件统计
   - 各类风险（监管、市场、运营）评估

# 数据输入

{json.dumps(data, ensure_ascii=False, indent=2)}

# 输出要求

请直接输出Markdown格式的报告内容，不要添加任何解释或前缀。报告应该完整、专业，适合呈现给高层管理人员。
"""

                    def _call_llm():
                        response = llm.invoke(prompt) if hasattr(llm, "invoke") else None
                        if hasattr(response, "content"):
                            return response.content
                        return str(response) if response is not None else ""

                    markdown = _invoke_with_timeout(_call_llm, float(llm_timeout_sec))
                    result = {"status": "success", "markdown": markdown}

                except Exception as exc:
                    # LLM 失败时使用 fallback
                    print(f"[WARN] LLM failed for report generation: {exc}")
                    import traceback
                    traceback.print_exc()

                    # 生成简单的 fallback 报告
                    fallback_lines = [
                        f"# {title}",
                        "",
                        "## 数据概览",
                        "",
                        "```json",
                        json.dumps(data, ensure_ascii=False, indent=2),
                        "```",
                        "",
                        "**注意**: 由于LLM服务异常，此报告为原始数据展示。",
                    ]
                    fallback = "\n".join(fallback_lines)
                    result = {
                        "status": "success",
                        "markdown": fallback,
                        "warning": f"llm_fallback: {type(exc).__name__}",
                    }

            # 如果提供了结构化的 sections，使用原有逻辑
            elif sections and isinstance(sections, list):
                use_llm = req.arguments.get("use_llm", True)
                if use_llm:
                    try:
                        llm = get_llm_by_type("basic")
                        prompt_lines = [
                            "你是企业分析报告写作者，请基于结构化输入生成Markdown报告。",
                            f"标题：{title}",
                            f"要求：{instruction}",
                            "结构化输入：",
                        ]
                        for sec in sections:
                            if not isinstance(sec, dict):
                                continue
                            heading = sec.get("heading", "")
                            content = sec.get("content", "")
                            if heading:
                                prompt_lines.append(f"- {heading}: {content}")
                        prompt = "\n".join(prompt_lines)

                        def _call_llm():
                            response = llm.invoke(prompt) if hasattr(llm, "invoke") else None
                            if hasattr(response, "content"):
                                return response.content
                            return str(response) if response is not None else ""

                        markdown = _invoke_with_timeout(_call_llm, float(llm_timeout_sec))
                        result = {"status": "success", "markdown": markdown}
                    except Exception as exc:
                        print(f"[WARN] LLM failed, using fallback: {exc}")
                        fallback = _build_markdown_report(title, sections)
                        result = {
                            "status": "success",
                            "markdown": fallback,
                            "warning": f"llm_fallback: {type(exc).__name__}",
                        }
                else:
                    result = {"status": "success", "markdown": _build_markdown_report(title, sections)}

            else:
                # 没有提供数据或 sections
                result = {
                    "status": "error",
                    "error": "Missing required parameter: 'data' or 'sections'"
                }

        except Exception as exc:
            import traceback
            error_msg = str(exc) if str(exc) else f"{type(exc).__name__}: {repr(exc)}"
            print(f"[ERROR] Report builder failed: {error_msg}")
            print(traceback.format_exc())
            result = {"status": "error", "error": error_msg, "traceback": traceback.format_exc()}
    elif req.tool == "remote_email_tool":
        print(f"[TOOL] remote_email_tool called")
        print(f"[TOOL] Request arguments: {json.dumps(req.arguments, ensure_ascii=False, indent=2)}")
        print(f"[TOOL] Arguments type: {type(req.arguments)}")
        print(f"[TOOL] Arguments keys: {list(req.arguments.keys()) if isinstance(req.arguments, dict) else 'N/A'}")

        try:
            data = _load_emails()
            emails = data.get("emails", [])

            # Extract parameters
            to_value = req.arguments.get("to")
            body_value = req.arguments.get("body", "")

            print(f"[TOOL] Extracted 'to': {to_value}")
            print(f"[TOOL] Extracted 'body' length: {len(str(body_value))}")

            payload = {
                "id": f"email-{len(emails)+1:04d}",
                "from": req.arguments.get("from", "noreply@internal.local"),
                "to": to_value,
                "subject": req.arguments.get("subject", "邮件发送"),
                "body": body_value,
                "attachments": req.arguments.get("attachments", []),
            }

            print(f"[TOOL] Created payload:")
            print(f"[TOOL]   id: {payload['id']}")
            print(f"[TOOL]   from: {payload['from']}")
            print(f"[TOOL]   to: {payload['to']}")
            print(f"[TOOL]   subject: {payload['subject']}")
            print(f"[TOOL]   body length: {len(str(payload['body']))}")

            emails.append(payload)
            data["emails"] = emails
            _EMAIL_CACHE = data
            _email_path().write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
            result = {"status": "success", "sent": payload}

            print(f"[TOOL] Email saved successfully")
        except Exception as exc:
            print(f"[TOOL] Error: {exc}")
            import traceback
            traceback.print_exc()
            result = {"status": "error", "error": str(exc)}
    elif req.tool == "remote_schedule_tool":
        try:
            data = _load_schedules()
            schedules = data.get("schedules", [])
            action = req.arguments.get("action", "create")
            if action == "create":
                payload = req.arguments.get("schedule") or {}
                if not isinstance(payload, dict):
                    raise ValueError("schedule must be an object")
                payload["id"] = payload.get("id") or f"visit-{len(schedules)+1:04d}"
                schedules.append(payload)
                data["schedules"] = schedules
                _SCHEDULE_CACHE = data
                _schedule_path().write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
                result = {"status": "success", "created": payload}
            else:
                start_date = req.arguments.get("start_date")
                end_date = req.arguments.get("end_date")
                filtered = []
                for item in schedules:
                    date = item.get("date")
                    if start_date and date and date < start_date:
                        continue
                    if end_date and date and date > end_date:
                        continue
                    filtered.append(item)
                result = {"status": "success", "matched_count": len(filtered), "schedules": filtered}
        except Exception as exc:
            result = {"status": "error", "error": str(exc)}
    elif req.tool == "remote_salary_info_tool":
        try:
            data = _load_salary()
            salary_records = data.get("salary_records", [])
            employee_id = req.arguments.get("employee_id")
            employee_name = req.arguments.get("employee_name")
            employee_id_list = req.arguments.get("employee_id_list")
            include_breakdown = req.arguments.get("include_breakdown", True)
            return_list = req.arguments.get("return_list", False)

            logger.info(f"Salary query - employee_id: {employee_id}, employee_name: {employee_name}, employee_id_list: {employee_id_list}")

            # 如果提供了 employee_id_list，批量查询
            if employee_id_list and isinstance(employee_id_list, list):
                matched_records = []
                for record in salary_records:
                    record_id = str(record.get("employee_id"))
                    if record_id in [str(eid) for eid in employee_id_list]:
                        salary_record = {
                            "status": "success",
                            "employee_id": record.get("employee_id"),
                            "employee_name": record.get("employee_name"),
                            "id_number": record.get("id_number"),
                            "monthly_salary": record.get("monthly_salary"),
                            "annual_salary": record.get("annual_salary"),
                            "currency": record.get("currency", "CNY"),
                            "last_updated": record.get("last_updated")
                        }
                        if include_breakdown:
                            salary_record["salary_breakdown"] = record.get("salary_breakdown", {})
                        matched_records.append(salary_record)

                logger.info(f"Batch salary query: found {len(matched_records)} records out of {len(employee_id_list)} requested")
                result = matched_records
            else:
                # 单个查询（原有逻辑）
                matched_record = None
                for record in salary_records:
                    if employee_id and str(record.get("employee_id")) == str(employee_id):
                        matched_record = record
                        break
                    if employee_name and record.get("employee_name") == employee_name:
                        matched_record = record
                        break

                if not matched_record:
                    result = {
                        "status": "not_found",
                        "error": f"No salary record found for employee_id={employee_id}, employee_name={employee_name}",
                        "searched_records": len(salary_records)
                    }
                else:
                    result = {
                        "status": "success",
                        "employee_id": matched_record.get("employee_id"),
                        "employee_name": matched_record.get("employee_name"),
                        "id_number": matched_record.get("id_number"),
                        "monthly_salary": matched_record.get("monthly_salary"),
                        "annual_salary": matched_record.get("annual_salary"),
                        "currency": matched_record.get("currency", "CNY"),
                        "last_updated": matched_record.get("last_updated")
                    }

                    if include_breakdown:
                        result["salary_breakdown"] = matched_record.get("salary_breakdown", {})

                    logger.info(f"Salary found: {result['employee_name']} - Monthly: {result['monthly_salary']}")

        except Exception as exc:
            import traceback
            logger.error(f"Salary query error: {exc}")
            logger.error(traceback.format_exc())
            result = {"status": "error", "error": str(exc)}
    elif req.tool == "remote_hr_complete_info_tool":
        try:
            # 这个工具同时查询人员信息和工资信息
            employee_name = req.arguments.get("employee_name")
            employee_id = req.arguments.get("employee_id")

            logger.info(f"HR complete info query - employee_name: {employee_name}, employee_id: {employee_id}")

            # 1. 查询人员信息
            person_sample = _load_sample()
            people = person_sample.get("personInfoList", [])

            matched_person = None
            for person in people:
                if employee_name and person.get("adtEmpeNm") == employee_name:
                    matched_person = person
                    break
                if employee_id and str(person.get("idvId")) == str(employee_id):
                    matched_person = person
                    break

            if not matched_person:
                result = {
                    "status": "not_found",
                    "error": f"No employee found for name={employee_name}, id={employee_id}"
                }
            else:
                # 2. 查询工资信息
                salary_data = _load_salary()
                salary_records = salary_data.get("salary_records", [])

                matched_salary = None
                person_name = matched_person.get("adtEmpeNm")
                person_id = matched_person.get("idvId")

                for record in salary_records:
                    if record.get("employee_name") == person_name:
                        matched_salary = record
                        break
                    if str(record.get("employee_id")) == str(person_id):
                        matched_salary = record
                        break

                # 3. 合并结果
                result = {
                    "status": "success",
                    "employee_info": {
                        "name": matched_person.get("adtEmpeNm"),
                        "id_number": matched_person.get("a001735"),  # 15位工号作为身份证号的替代
                        "employee_id": matched_person.get("idvId"),
                        "gender": matched_person.get("gnd"),
                        "birth_date": matched_person.get("brthDt"),
                        "age": matched_person.get("age"),
                        "position": matched_person.get("tcoPostNm"),
                        "department": matched_person.get("holdposInstNm"),
                        "join_date": matched_person.get("jnCcbDt"),
                        "work_years": matched_person.get("pcsTrdYrlmt"),
                        "education": matched_person.get("hgstEddgrNm"),
                        "phone": matched_person.get("officePhone"),
                        "email": matched_person.get("internalMaiBox")
                    }
                }

                if matched_salary:
                    result["salary_info"] = {
                        "monthly_salary": matched_salary.get("monthly_salary"),
                        "annual_salary": matched_salary.get("annual_salary"),
                        "currency": matched_salary.get("currency", "CNY"),
                        "salary_breakdown": matched_salary.get("salary_breakdown", {})
                    }
                    logger.info(f"Complete info found: {person_name} - Monthly: {matched_salary.get('monthly_salary')}")
                else:
                    result["salary_info"] = None
                    logger.warning(f"Salary info not found for {person_name}")

        except Exception as exc:
            import traceback
            logger.error(f"HR complete info query error: {exc}")
            logger.error(traceback.format_exc())
            result = {"status": "error", "error": str(exc)}
    elif req.tool == "remote_docx_generator_tool":
        try:
            import datetime
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

            template_name = req.arguments.get("template_name")
            data = req.arguments.get("data", {})
            output_filename = req.arguments.get("output_filename") or f"document_{int(time.time())}"

            logger.info(f"Generating document: template={template_name}, filename={output_filename}")

            # Load template configuration
            templates = _load_templates()
            template_config = templates.get("templates", {}).get(template_name)

            if not template_config:
                raise ValueError(f"Template '{template_name}' not found")

            # Prepare data with defaults
            current_date = datetime.datetime.now().strftime("%Y年%m月%d日")
            company_name = data.get("company_name", "中国建设银行股份有限公司")

            # Convert salary to Chinese if needed
            if "monthly_salary" in data and "monthly_salary_cn" not in data:
                data["monthly_salary_cn"] = _number_to_chinese(float(data["monthly_salary"]))
            if "annual_salary" in data and "annual_salary_cn" not in data:
                data["annual_salary_cn"] = _number_to_chinese(float(data["annual_salary"]))

            data["current_date"] = current_date
            data["company_name"] = company_name

            # Create document
            doc = Document()

            # Set default font
            style = doc.styles['Normal']
            style.font.name = '宋体'
            style.font.size = Pt(12)

            # Generate content from template
            for item in template_config.get("content", []):
                item_type = item.get("type")
                text = item.get("text", "")

                # Replace placeholders
                for key, value in data.items():
                    placeholder = "{{" + key + "}}"
                    if placeholder in text:
                        text = text.replace(placeholder, str(value))

                if item_type == "heading":
                    level = item.get("level", 1)
                    para = doc.add_heading(text, level=level)
                    align = item.get("align", "left")
                    if align == "center":
                        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    elif align == "right":
                        para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                elif item_type == "paragraph":
                    para = doc.add_paragraph(text)
                    align = item.get("align", "left")
                    if align == "center":
                        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    elif align == "right":
                        para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                    indent = item.get("indent", 0)
                    if indent > 0:
                        para.paragraph_format.first_line_indent = Inches(indent * 0.5)

            # Save document
            output_dir = Path(__file__).resolve().parent / "output"
            output_dir.mkdir(exist_ok=True)
            output_path = output_dir / f"{output_filename}.docx"

            doc.save(str(output_path))

            logger.info(f"Document generated successfully: {output_path}")

            result = {
                "status": "success",
                "file_path": str(output_path),
                "file_name": f"{output_filename}.docx",
                "template_used": template_name,
                "message": f"文档已生成：{output_filename}.docx"
            }

        except Exception as exc:
            import traceback
            logger.error(f"Document generation error: {exc}")
            logger.error(traceback.format_exc())
            result = {"status": "error", "error": str(exc), "traceback": traceback.format_exc()}
    elif req.tool == "knowledge_search_tool":
        try:
            data = _load_knowledge()
            knowledge_items = data.get("knowledge_items", [])
            query = req.arguments.get("query", "")

            if not query:
                raise ValueError("query parameter is required")

            # 使用LLM理解问题并从知识库中检索答案
            llm = get_llm_by_type("basic")

            # 构建知识库内容字符串
            knowledge_content = []
            for idx, item in enumerate(knowledge_items, 1):
                knowledge_content.append(f"[知识条目 {idx}]")
                knowledge_content.append(f"类别: {item.get('category', '')}")
                knowledge_content.append(f"问题: {item.get('question', '')}")
                knowledge_content.append(f"内容:\n{item.get('content', '')}")
                knowledge_content.append("")

            knowledge_text = "\n".join(knowledge_content)

            # 构建prompt
            prompt = f"""你是一位专业的知识库查询助手。用户提出了一个问题，请从以下知识库中找到相关信息，并生成详细的回答。

# 用户问题
{query}

# 知识库内容
{knowledge_text}

# 回答要求
1. 仔细分析用户问题，理解其真实意图
2. 从知识库中找到最相关的信息
3. 基于知识库内容生成详细、准确的回答
4. 如果知识库中有法律法规依据，请在回答中引用
5. 回答应该完整、专业，包含必要的说明和注意事项
6. 如果知识库中没有相关信息，请明确告知用户

请直接输出回答内容，不要添加任何前缀或解释。"""

            # 调用LLM
            response = llm.invoke(prompt) if hasattr(llm, "invoke") else None
            if hasattr(response, "content"):
                answer = response.content
            else:
                answer = str(response) if response is not None else "无法生成回答"

            # 记录完整答案
            logger.info(f"Knowledge search completed for query: {query}")
            logger.info(f"Full answer:\n{answer}")

            result = {
                "status": "success",
                "query": query,
                "answer": answer,
                "knowledge_items_count": len(knowledge_items)
            }
        except Exception as exc:
            import traceback
            result = {
                "status": "error",
                "error": str(exc),
                "traceback": traceback.format_exc()
            }
    else:
        result = f"[remote-tool:{req.tool}] ok"

    return {
        "status": "success",
        "result": result,
        "metadata": {"has_auth": bool(authorization)},
    }


@app.post("/skill")
async def skill(req: SkillRequest, authorization: Optional[str] = Header(default=None)):
    if req.skill == "remote_summarize":
        text = req.arguments.get("text", "")
        result = f"[remote-skill] summary: {text[:30]}"
    else:
        result = f"[remote-skill:{req.skill}] ok"
    return {
        "status": "success",
        "result": result,
        "metadata": {"has_auth": bool(authorization)},
    }


if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="0.0.0.0", port=8011)
